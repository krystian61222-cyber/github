"""Part 1: Setup and sections 1-3 of the GitHub case study."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_doc():
    doc = Document()
    return doc


def add_italic_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    return p


def add_bold_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text, style='Normal')
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p


def add_table_row_bold_first(table, row_idx, cells_data):
    row = table.rows[row_idx]
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        if i == 0:
            run.bold = True


def build_sections_1_to_3(doc):
    # Header lines
    add_italic_para(doc, 'Fundusze Venture Capital i Private Equity \u2013 Wariant B')
    add_italic_para(doc, 'Andreessen Horowitz (a16z)')

    # Title
    doc.add_heading('Case study: Inwestycja a16z w GitHub Inc.', level=0)

    # Section 1
    doc.add_heading('1. T\u0142o i charakterystyka sp\u00f3\u0142ki', level=1)

    add_body(doc,
        'GitHub Inc. to ameryka\u0144ska sp\u00f3\u0142ka technologiczna za\u0142o\u017cona w 2008 r. przez Toma Prestona-Wernera, '
        'Chrisa Wanstratha i PJ Hyetta z siedzib\u0105 w San Francisco (Kalifornia). Sp\u00f3\u0142ka prowadzi platform\u0119 '
        'hostingu repozytori\u00f3w kodu \u017ar\u00f3d\u0142owego opart\u0105 na systemie kontroli wersji Git, stworzon\u0105 '
        'przez Linusa Torvaldsa w 2005 r. GitHub umo\u017cliwia programistom przechowywanie, wersjonowanie i '
        'wsp\u00f3\u0142dzielenie kodu, a tak\u017ce wsp\u00f3\u0142prac\u0119 nad projektami open source i komercyjnymi '
        'za pomoc\u0105 mechanizm\u00f3w pull request, code review, issue tracking i wikis.')

    add_body(doc,
        'Model biznesowy GitHub opiera\u0142 si\u0119 na modelu freemium: darmowe repozytoria publiczne dla '
        'projekt\u00f3w open source oraz p\u0142atne repozytoria prywatne dla firm i indywidualnych programist\u00f3w. '
        'W 2011 r. sp\u00f3\u0142ka wprowadzi\u0142a GitHub Enterprise \u2013 wersj\u0119 on-premise przeznaczon\u0105 '
        'dla korporacji wymagaj\u0105cych hostingu kodu za w\u0142asnym firewallem. Model przychodowy obejmowa\u0142 '
        'trzy filary: (1) subskrypcje indywidualne i zespo\u0142owe na prywatne repozytoria, (2) GitHub Enterprise '
        'Server (instalacja lokalna) oraz (3) GitHub Enterprise Cloud (hosted). Sp\u00f3\u0142ka by\u0142a bootstrapowana '
        'od za\u0142o\u017cenia w 2008 r. do momentu inwestycji a16z w 2012 r. \u2013 jeden z najd\u0142u\u017cszych '
        'okres\u00f3w samofinansowania w historii Silicon Valley dla firmy o takiej skali.')

    add_body(doc,
        'W momencie wej\u015bcia a16z (lipiec 2012) GitHub obs\u0142ugiwa\u0142 ponad 3,5 miliona u\u017cytkownik\u00f3w '
        'i miliony repozytori\u00f3w kodu, b\u0119d\u0105c de facto standardem bran\u017cowym dla hosting\u00f3w Git. '
        'Platforma pe\u0142ni\u0142a unikaln\u0105 rol\u0119 \u201csieci spo\u0142eczno\u015bciowej dla kodu\u201d \u2013 '
        'programi\u015bci \u015bledzili projekty i innych developer\u00f3w, a aktywno\u015b\u0107 na GitHubie stawa\u0142a '
        'si\u0119 de facto CV technologicznym. Efekty sieciowe by\u0142y silne: im wi\u0119cej projekt\u00f3w open source '
        'hostowano na GitHubie, tym wi\u0119cej programist\u00f3w do\u0142\u0105cza\u0142o, co przyci\u0105ga\u0142o '
        'kolejne projekty \u2013 klasyczny flywheel effect. Konkurencja obejmowa\u0142a Bitbucket (Atlassian, darmowe '
        'repozytoria prywatne), GitLab (open source, zintegrowane CI/CD) oraz SourceForge (w zaniku).')

    # Section 2
    doc.add_heading('2. Parametry transakcji', level=1)

    table = doc.add_table(rows=8, cols=2, style='Table Grid')
    data = [
        ('Runda finansowania', 'Series A'),
        ('Data zamkni\u0119cia', 'lipiec 2012 r.'),
        ('Warto\u015b\u0107 rundy', '100 mln USD'),
        ('Lead investor', 'Andreessen Horowitz (jedyny inwestor \u2013 sole investor)'),
        ('Wycena post-money', 'ok. 750 mln USD'),
        ('Miejsce w zarz\u0105dzie', 'Peter Levine (General Partner a16z) do\u0142\u0105cza do Board of Directors'),
        ('Bran\u017ca', 'Developer tools / infrastruktura oprogramowania / SaaS'),
        ('\u0141\u0105czne finansowanie do rundy', '0 USD (sp\u00f3\u0142ka bootstrapowana od 2008 r.)'),
    ]
    for i, (param, val) in enumerate(data):
        add_table_row_bold_first(table, i, [param, val])

    add_body(doc,
        'Inwestycja a16z w GitHub by\u0142a jedn\u0105 z najwi\u0119kszych rund Series A w historii Silicon Valley '
        'w momencie jej zamkni\u0119cia. Wyj\u0105tkowo\u015b\u0107 transakcji polega\u0142a na trzech elementach: '
        '(1) 100 mln USD od jednego inwestora w sp\u00f3\u0142k\u0119 bez wcze\u015bniejszego finansowania VC, '
        '(2) wycena 750 mln USD przy braku przychod\u00f3w typowych dla sp\u00f3\u0142ki SaaS o takiej warto\u015bci, '
        '(3) ca\u0142kowity brak wsp\u00f3\u0142inwestor\u00f3w \u2013 a16z przej\u0105\u0142 ca\u0142o\u015b\u0107 rundy.')

    # Section 3
    doc.add_heading('3. Cel inwestycji i teza inwestycyjna a16z', level=1)

    add_body(doc,
        'Inwestycja zosta\u0142a zrealizowana z g\u0142\u00f3wnego funduszu venture a16z (Fund III). '
        '\u015arodki mia\u0142y sfinansowa\u0107 rozbudow\u0119 zespo\u0142u (engineering, sprzeda\u017c enterprise, '
        'operations), rozw\u00f3j produktu GitHub Enterprise oraz przyspieszenie ekspansji korporacyjnej. '
        'By\u0142a to jedna z najwi\u0119kszych pojedynczych rund Series A w historii \u2013 100 mln USD od jednego '
        'inwestora w sp\u00f3\u0142k\u0119, kt\u00f3ra nigdy wcze\u015bniej nie przyj\u0119\u0142a kapita\u0142u VC.')

    add_body(doc,
        'Teza inwestycyjna Petera Levine\u2019a i a16z opiera\u0142a si\u0119 na pi\u0119ciu filarach:')

    add_bullet(doc,
        'Narz\u0119dzia programistyczne jako infrastruktura krytyczna \u2013 Git sta\u0142 si\u0119 standardem '
        'kontroli wersji, a GitHub jedynym dominuj\u0105cym hostem. Ka\u017cda firma tworz\u0105ca oprogramowanie '
        'potrzebowa\u0142a repozytorium kodu, co czyni\u0142o GitHub infrastruktur\u0105 o charakterze '
        'quasi-monopolistycznym.')

    add_bullet(doc,
        '\u201eSoftware eating the world\u201d zastosowane do narz\u0119dzi developerskich \u2013 manifest '
        'Marca Andreessena z sierpnia 2011 r. (WSJ, \u201eWhy Software Is Eating the World\u201d) stanowi\u0142 '
        'intelektualn\u0105 ram\u0119 dla tezy, \u017ce skoro oprogramowanie zjada ka\u017cd\u0105 bran\u017c\u0119, '
        'to narz\u0119dzia do jego tworzenia s\u0105 \u201ekilofami w gor\u0105czce z\u0142ota\u201d \u2013 '
        'infrastruktur\u0105 korzystaj\u0105c\u0105 na wzro\u015bcie ca\u0142ego ekosystemu.')

    add_bullet(doc,
        'GitHub jako sie\u0107 spo\u0142eczno\u015bciowa dla kodu z efektami sieciowymi \u2013 profil '
        'programisty na GitHubie stawa\u0142 si\u0119 CV technologicznym, a aktywno\u015b\u0107 w projektach '
        'open source budowa\u0142a reputacj\u0119. Im wi\u0119cej projekt\u00f3w, tym wi\u0119cej developer\u00f3w, '
        'tym wi\u0119cej projekt\u00f3w \u2013 klasyczny network effect.')

    add_bullet(doc,
        'Potencja\u0142 ekspansji enterprise \u2013 GitHub Enterprise (on-premise) umo\u017cliwia\u0142 dost\u0119p '
        'do bud\u017cet\u00f3w IT korporacji (Google, Microsoft, Facebook, Amazon), kt\u00f3re potrzebowa\u0142y '
        'bezpiecznego hostingu kodu za firewallem. Levine widzia\u0142 analogie do modelu Red Hat \u2013 '
        'open-source community jako engine adopcji, enterprise jako engine przychod\u00f3w.')

    add_bullet(doc,
        'Samonaprawiaj\u0105ce si\u0119 efekty spo\u0142eczno\u015bciowe open source \u2013 spo\u0142eczno\u015b\u0107 '
        'open source generowa\u0142a warto\u015b\u0107 bezp\u0142atnie (miliony projekt\u00f3w, dokumentacja, '
        'integracje), a GitHub czerpa\u0142 korzy\u015bci jako platforma hostuj\u0105ca t\u0119 aktywno\u015b\u0107. '
        'Koszt pozyskania u\u017cytkownika (CAC) by\u0142 bliski zeru dla segmentu open source.')

    add_body(doc,
        'Peter Levine we wpisie blogowym a16z \u201eSoftware Eating Software Development\u201d (9 lipca 2012 r.) '
        'uj\u0105\u0142 tez\u0119 nast\u0119puj\u0105co: je\u015bli oprogramowanie zjada \u015bwiat, to narz\u0119dzia '
        'do jego tworzenia s\u0105 fundamentem ca\u0142ej transformacji cyfrowej. GitHub \u2013 jako dominuj\u0105ca '
        'platforma wsp\u00f3\u0142pracy programist\u00f3w \u2013 zajmuje pozycj\u0119 analogiczn\u0105 do tej, '
        'kt\u00f3r\u0105 Amazon Web Services zajmuje w infrastrukturze chmurowej: warstwa bazowa, od kt\u00f3rej '
        'zale\u017cy ca\u0142y ekosystem.')

    return doc
