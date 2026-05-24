"""Part 2: Sections 4-5 of the GitHub case study."""


def build_sections_4_to_5(doc):
    # Section 4
    doc.add_heading('4. Osoba odpowiedzialna za transakcj\u0119', level=1)

    from docx.shared import Pt

    p = doc.add_paragraph()
    run = p.add_run('Peter Levine')
    run.bold = True
    p.add_run(
        ' \u2013 General Partner w Andreessen Horowitz (od 2010 r.), prowadzi\u0142 transakcj\u0119, '
        'podpisa\u0142 termsheet i obj\u0105\u0142 miejsce w radzie dyrektor\u00f3w GitHub po zamkni\u0119ciu rundy. '
        'Przed do\u0142\u0105czeniem do a16z Levine by\u0142 CEO XenSource \u2013 firmy tworz\u0105cej '
        'open-source\u2019ow\u0105 platform\u0119 wirtualizacji, kt\u00f3r\u0105 sprzeda\u0142 Citrix Systems '
        'w 2007 r. za oko\u0142o 500 mln USD (w momencie przej\u0119cia). Nast\u0119pnie pe\u0142ni\u0142 '
        'funkcj\u0119 VP i GM w Citrix, zarz\u0105dzaj\u0105c pionem wirtualizacji. Posiada tytu\u0142 '
        'doktora (PhD) MIT Sloan School of Management.'
    )

    p2 = doc.add_paragraph(
        'Levine specjalizowa\u0142 si\u0119 w inwestycjach w infrastruktur\u0119 oprogramowania, '
        'narz\u0119dzia programistyczne i systemy enterprise. Jego do\u015bwiadczenie operacyjne '
        'w XenSource \u2013 firmie, kt\u00f3ra skomercjalizowa\u0142a technologi\u0119 open source '
        '(Xen hypervisor) i sprzeda\u0142a j\u0105 korporacji \u2013 czyni\u0142o go idealnym '
        'partnerem dla GitHub, kt\u00f3ry realizowa\u0142 analogiczny model: open-source community '
        'jako silnik adopcji, enterprise jako silnik przychod\u00f3w.'
    )

    p3 = doc.add_paragraph(
        'W kontek\u015bcie GitHub Levine pe\u0142ni\u0142 aktywn\u0105 rol\u0119 cz\u0142onka rady '
        'dyrektor\u00f3w, doradzaj\u0105c w kwestiach strategii enterprise, rekrutacji kadry '
        'zarz\u0105dzaj\u0105cej oraz ekspansji produktowej. Jego kluczowe publikacje zwi\u0105zane '
        'z tez\u0105 inwestycyjn\u0105 obejmuj\u0105:'
    )

    from gen_part1 import add_bullet
    add_bullet(doc,
        '\u201eSoftware Eating Software Development\u201d \u2013 wpis blogowy a16z, 9 lipca 2012 r. '
        '\u2013 og\u0142oszenie inwestycji i sformu\u0142owanie tezy o narz\u0119dziach developerskich '
        'jako fundamentalnej warstwie transformacji cyfrowej.')
    add_bullet(doc,
        '\u201eWhy There Will Never Be Another Red Hat: The Economics of Open Source\u201d \u2013 '
        'esej a16z, 2014 r. \u2013 analiza modeli komercjalizacji open source, bezpo\u015brednio '
        'relewantna dla strategii GitHub Enterprise.')
    add_bullet(doc,
        '\u201eOpen Source: From Community to Commercialization\u201d \u2013 wyk\u0142ad Stanford/a16z, '
        '2016 r. \u2013 rozbudowana analiza ekonomiki open source z przyk\u0142adami m.in. GitHub, '
        'Docker, Cloudera.')

    p4 = doc.add_paragraph(
        'Wsp\u00f3\u0142sygnatariuszami inwestycji byli Marc Andreessen i Ben Horowitz '
        '(wsp\u00f3\u0142za\u0142o\u017cyciele funduszu), kt\u00f3rzy aktywnie uczestniczyli '
        'w procesie due diligence i p\u00f3\u017aniej wspierali sp\u00f3\u0142k\u0119 w momentach '
        'kryzysowych (m.in. zmiana CEO w 2014 r.).'
    )

    # Section 5
    doc.add_heading('5. Dzia\u0142ania a16z po inwestycji', level=1)

    p5 = doc.add_paragraph(
        'Wsparcie a16z dla GitHub po wej\u015bciu kapita\u0142owym mia\u0142o charakter zar\u00f3wno '
        'strategiczny, jak i operacyjny \u2013 typowy dla inwestycji wczesno-wzrostowych funduszu '
        'z silnym zapleczem operacyjnym (tzw. platform model). A16z zaanga\u017cowa\u0142 pe\u0142n\u0105 '
        'infrastruktur\u0119 operacyjn\u0105 funduszu w skalowanie GitHub:'
    )

    add_bullet(doc,
        'Pozyskanie talentu mened\u017cerskiego i rekrutacja \u2013 a16z pom\u00f3g\u0142 zbudowa\u0107 '
        'ponad 70-osobowy wewn\u0119trzny zesp\u00f3\u0142 rekrutacyjny GitHub oraz obsadzi\u0107 '
        'kluczowe role lider\u00f3w sprzeda\u017cy enterprise. Fundusz wykorzysta\u0142 w\u0142asny '
        'executive talent team do identyfikacji i pozyskania kandydat\u00f3w z do\u015bwiadczeniem '
        'w sprzeda\u017cy B2B (m.in. z SAP, Oracle, Salesforce). To jeden z najbardziej wymiernych '
        'przyk\u0142ad\u00f3w operacyjnego wsparcia a16z dla sp\u00f3\u0142ki portfelowej.')

    add_bullet(doc,
        'Strategia enterprise go-to-market \u2013 Levine, bazuj\u0105c na do\u015bwiadczeniu '
        'z XenSource i Citrix, pom\u00f3g\u0142 zdefiniowa\u0107 model sprzeda\u017cy enterprise: '
        'segmentacj\u0119 klient\u00f3w (SMB vs. mid-market vs. Fortune 500), pricing GitHub Enterprise '
        '(per-seat vs. per-server), procesy sprzeda\u017cy bezpo\u015bredniej (direct sales) oraz '
        'programy partnerskie. To by\u0142a kluczowa transformacja \u2013 z firmy narz\u0119dziowej '
        'dla developer\u00f3w w platform\u0119 enterprise software.')

    add_bullet(doc,
        'Skalowanie organizacji in\u017cynierskiej \u2013 Levine dzieli\u0142 si\u0119 '
        'do\u015bwiadczeniem operacyjnym z budowy i skalowania zespo\u0142u XenSource/Citrix '
        '(od startup do kilkuset os\u00f3b), doradzaj\u0105c w kwestiach struktury organizacyjnej, '
        'proces\u00f3w engineeringowych i kultury technicznej w fazie hiperwzrostu.')

    add_bullet(doc,
        'Budowanie marki i PR \u2013 a16z konsekwentnie pozycjonowa\u0142 GitHub jako '
        '\u201ekrytyczn\u0105 infrastruktur\u0119 enterprise\u201d, a nie jedynie narz\u0119dzie '
        'dla spo\u0142eczno\u015bci open source. Publikacje Levine\u2019a (\u201eSoftware Eating '
        'Software Development\u201d, \u201eWhy There Will Never Be Another Red Hat\u201d) budowa\u0142y '
        'narracyjn\u0105 ram\u0119 uzasadniaj\u0105c\u0105 wydatki IT korporacji na GitHub Enterprise.')

    add_bullet(doc,
        'Wprowadzenia strategiczne do Fortune 500 \u2013 a16z wykorzysta\u0142 sie\u0107 kontakt\u00f3w '
        'funduszu do organizowania spotka\u0144 z decydentami IT w najwi\u0119kszych korporacjach '
        '(Google, Microsoft, Facebook, Amazon, banki inwestycyjne). Efekt: masowa adopcja GitHub '
        'Enterprise przez Fortune 500 w latach 2013\u20132017.')

    add_bullet(doc,
        'Wsparcie w kryzysie CEO (2014) \u2013 gdy Tom Preston-Werner zrezygnowa\u0142 w kwietniu '
        '2014 r. w wyniku kontrowersji (dochodzenie wewn\u0119trzne dot. mobbingu \u2013 opisane '
        'przez Wired jako \u201eThe GitHub Investigation\u201d), a16z aktywnie wspiera\u0142 proces '
        'tranzycji do nowego CEO (Chris Wanstrath). Ben Horowitz \u2013 autor ksi\u0105\u017cki '
        '\u201eThe Hard Thing About Hard Things\u201d o zarz\u0105dzaniu w kryzysie \u2013 pe\u0142ni\u0142 '
        'rol\u0119 mentora w tym okresie.')

    p6 = doc.add_paragraph(
        'Model wsparcia a16z dla GitHub mo\u017cna por\u00f3wna\u0107 z analogicznym zaanga\u017cowaniem '
        'w innych sp\u00f3\u0142kach portfelowych: podobnie jak w Lyft (rekrutacja, public policy, '
        'mentoring boardowy), ale z silniejszym akcentem enterprise \u2013 analogicznie do p\u00f3\u017aniejszej '
        'pracy z Coinbase nad pozycjonowaniem regulacyjnym w sektorze kryptowalut.'
    )

    return doc
